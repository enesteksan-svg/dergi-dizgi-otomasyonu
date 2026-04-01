import streamlit as st
from docx import Document
from weasyprint import HTML
import io
import re

# --- ANALİZ FONKSİYONLARI ---
def analiz_et(doc):
    tam_metin = ""
    kaynaklar = []
    kaynakca_basladi = False
    hatalar = []
    
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if t.lower() in ["kaynakça", "references", "kaynaklar"]:
            kaynakca_basladi = True
            continue
        if kaynakca_basladi:
            kaynaklar.append(t)
        else:
            tam_metin += " " + t
            
    # Atıf Denetimi
    metin_atiflari = re.findall(r'\(([A-ZÇĞİÖŞÜ][a-zçğıöşü]+,\s+\d{4})\)', tam_metin)
    kaynakca_blogu = " ".join(kaynaklar).lower()
    for atif in set(metin_atiflari):
        soyad = atif.split(',')[0].strip().lower()
        if soyad not in kaynakca_blogu:
            hatalar.append(f"⚠️ Atıf ({atif}) var ama kaynakçada '{soyad.capitalize()}' yok.")
            
    return hatalar, kaynaklar

# --- DİZGİ MOTORU ---
def dizgi_yap(doc, dergi_adi):
    html_sections = {"baslik": f"<h1>{dergi_adi}</h1>", "ozet": "", "metin": "", "kaynakca": ""}
    kaynakca_basladi = False
    ozet_basladi = False

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if t.lower() in ["özet", "abstract"]: ozet_basladi = True; continue
        if t.lower() in ["kaynakça", "references"]: kaynakca_basladi = True; continue
        
        if kaynakca_basladi: html_sections["kaynakca"] += f"<li>{t}</li>"
        elif ozet_basladi: html_sections["ozet"] += f"<p>{t}</p>"
        else: html_sections["metin"] += f"<p>{t}</p>"

    css = """
    @page { size: A4; margin: 2cm; @bottom-center { content: counter(page); } }
    body { font-family: 'Times New Roman', serif; text-align: justify; }
    h1 { text-align: center; text-transform: uppercase; border-bottom: 2px solid #000; }
    .ozet { font-style: italic; background: #f9f9f9; padding: 15px; margin-bottom: 20px; border-left: 5px solid #333; }
    .metin { column-count: 2; column-gap: 1cm; }
    .kaynakca { font-size: 10pt; margin-top: 30px; border-top: 1px solid #000; }
    """
    
    full_html = f"<html><style>{css}</style><body>{html_sections['baslik']}<div class='ozet'>{html_sections['ozet']}</div><div class='metin'>{html_sections['metin']}</div><div class='kaynakca'><h3>KAYNAKÇA</h3><ul>{html_sections['kaynakca']}</ul></div></body></html>"
    return full_html

# --- ARAYÜZ ---
st.title("📚 Dergi Dizgi Sistemi")
dergi_adi = st.text_input("Dergi Adı", "Akademik Bülten")
file = st.file_uploader("Word Dosyası Yükle", type="docx")

if file:
    doc = Document(file)
    if st.button("Dizgiyi ve Denetimi Başlat"):
        hatalar, _ = analiz_et(doc)
        if hatalar:
            st.warning("Bazı eksikler bulundu:")
            for h in hatalar: st.write(h)
        
        html_cikti = dizgi_yap(doc, dergi_adi)
        pdf_data = io.BytesIO()
        HTML(string=html_cikti).write_pdf(pdf_data)
        
        st.success("Dizgi Tamam!")
        st.download_button("📥 PDF İndir", pdf_data.getvalue(), "dizgi.pdf")